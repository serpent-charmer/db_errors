# -*- coding: utf-8 -*-
from docx import Document
from random import choice, sample
import sys

try:
    document = Document('db.docx')
except:
    print('Не найден переименованный db.docx')
    sys.exit(1)

document_out = Document()

q_amt = 95

idx = 0
error_percentage = int(input('Введите процент ошибок: ') or 15)
error_amount = int((q_amt * error_percentage / 100))
error_seq = sample(range(q_amt+1), error_amount)

report = ''

for p in document.paragraphs:
    _p = document_out.add_paragraph()
    _p.alignment = p.alignment
    ppf = p.paragraph_format
    _ppf = _p.paragraph_format
    _ppf.first_line_indent = ppf.first_line_indent
    _ppf.keep_together = ppf.keep_together
    _ppf.keep_with_next = ppf.keep_with_next
    _ppf.left_indent = ppf.left_indent
    _ppf.line_spacing = ppf.line_spacing
    _ppf.line_spacing_rule = ppf.line_spacing_rule
    _ppf.page_break_before = ppf.page_break_before
    _ppf.right_indent = ppf.right_indent
    _ppf.space_after = 0
    _ppf.space_before = ppf.space_before
    _ppf.widow_control = ppf.widow_control 
    _p.style = p.style
    
    for run in p.runs:
        output_run = _p.add_run() 
        output_run.bold = run.bold
        output_run.italic = run.italic
        output_run.underline = run.underline
        output_run.font.color.rgb = run.font.color.rgb
        if 'бал' in run.text:
            num = int(run.text.split('.')[0])
            output_run.text = run.text.split(')')[1]
            
            if not num in error_seq:
                continue
            
            opt = document.paragraphs[idx].text
            
            report += ('*Делаем ошибку в %d:\n%s...' % (num, opt[:70].rstrip(' '))) + '\n'
            
            try:
                parts = 0
                answer = 0
                for i in range(idx+1, idx+10):
                    pt = document.paragraphs[i].text
                    if not 'бал' in pt:
                        
                        
                        if '*' in pt:
                            answer = i
                            
                        if pt != '':
                            parts += 1
                            
                    else:
                        
                        break
                    
                document.paragraphs[answer].text = document.paragraphs[answer].text.split(' *')[0]
                
                poss = list(range(idx+1, idx+parts+1))
                poss.remove(answer)
                
                replaced = choice(poss)
                
                sp = document.paragraphs[replaced]
                sp.runs[len(sp.runs)-1].text += ' *'
                        
            except  Exception as e:
                pass
            
            continue
        output_run.text = run.text
        
    idx += 1
document_out.save('out.docx')
print(report)